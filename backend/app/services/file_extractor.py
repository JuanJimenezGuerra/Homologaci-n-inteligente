import PyPDF2
import docx
import pandas as pd
import io
import os
from thefuzz import fuzz, process
from sqlalchemy.orm import Session
from ..models import Cargo

SIMILARITY_THRESHOLD = 70  # minimum fuzzy match score (0-100)

def extract_text_from_pdf(file_bytes):
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(file_bytes):
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)

def _find_best_cargo_match(cargo_nombre: str, upload_id: int, db: Session):
    """Find existing cargo with name similar to cargo_nombre via fuzzy matching."""
    all_cargos = db.query(Cargo).filter(
        Cargo.upload_id == upload_id,
        Cargo.area != 'DESCRIPCION_ANEXA'
    ).all()
    if not all_cargos:
        return None, 0
    names = {c.id: c.nombre_cargo for c in all_cargos}
    best = process.extractOne(cargo_nombre, list(names.values()), scorer=fuzz.token_sort_ratio)
    if best and best[1] >= SIMILARITY_THRESHOLD:
        matched_id = [k for k, v in names.items() if v == best[0]][0]
        return matched_id, best[1]
    return None, 0

def process_extra_descriptions(upload_id: int, files: list, db: Session):
    """
    Processes files (PDF, DOCX, XLSX) and associates their text with existing cargos
    via fuzzy name matching. If a close match is found, the description is added to
    that cargo and marked as 'tiene_descripcion_anexa'. If no match, a warning is logged
    and NO separate cargo is created (instead, a placeholder cargo is created for review).
    """
    mapped_count = 0
    no_match_files = []

    for file_obj in files:
        filename = file_obj.filename
        content = file_obj.file.read()

        ext = os.path.splitext(filename)[1].lower()
        if ext not in ['.pdf', '.docx', '.doc', '.xlsx', '.xls']:
            continue

        try:
            text = ""
            if ext == '.pdf':
                text = extract_text_from_pdf(content)
            elif ext in ['.docx', '.doc']:
                text = extract_text_from_docx(content)
            elif ext in ['.xlsx', '.xls']:
                df = pd.read_excel(io.BytesIO(content))
                text = df.to_string()

            cargo_nombre = os.path.splitext(filename)[0].strip()
            if not cargo_nombre:
                continue

            if not text or not text.strip():
                text = f"Descripción del cargo: {cargo_nombre}"

            # Try fuzzy match against existing cargos
            matched_id, score = _find_best_cargo_match(cargo_nombre, upload_id, db)

            if matched_id:
                existing = db.query(Cargo).filter(Cargo.id == matched_id).first()
                if existing:
                    existing.descripcion_empresa = text
                    if not existing.area or existing.area == 'PENDIENTE':
                        existing.area = 'DESCRIPCION_ANEXA'
                    print(f"Matched '{cargo_nombre}' → existing cargo #{existing.id} '{existing.nombre_cargo}' (score={score}%)")
                    mapped_count += 1
            else:
                # No close match — create cargo with area='SIN_MATCH' for user review
                print(f"No match for '{cargo_nombre}' (best score={score}%) — created as SIN_MATCH")
                new_cargo = Cargo(
                    upload_id=upload_id,
                    nombre_cargo=cargo_nombre,
                    area='SIN_MATCH',
                    descripcion_empresa=text,
                    estado='PENDIENTE'
                )
                db.add(new_cargo)
                no_match_files.append(cargo_nombre)
                mapped_count += 1

        except Exception as e:
            print(f"Error processing file {filename}: {e}")
            import traceback
            traceback.print_exc()
            continue

    db.commit()
    print(f"Total extra description files processed: {mapped_count}")
    if no_match_files:
        print(f"Files with no matching cargo (created as SIN_MATCH): {no_match_files}")
    return mapped_count
